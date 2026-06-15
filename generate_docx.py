import subprocess
import sys
import os
import math

# Step 1: Programmatic installation of python-docx and pillow if missing
try:
    import docx
except ImportError:
    print("python-docx not found. Installing now...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    print("Pillow not found. Installing now...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "Pillow"])
    from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Color definitions
DEEP_BLUE = (26, 54, 93)
LIGHT_BLUE = (66, 153, 225)
BORDER_GRAY = (203, 213, 224)
TEXT_DARK = (45, 55, 72)
TEXT_LIGHT = (113, 128, 150)
BG_LIGHT = (247, 250, 252)

# Helper function to draw box
def draw_box(draw, coords, text, title="", fill_color=(235, 244, 255), border_color=(66, 153, 225), font_title=None, font_text=None):
    try:
        draw.rounded_rectangle(coords, radius=10, fill=fill_color, outline=border_color, width=3)
    except AttributeError:
        draw.rectangle(coords, fill=fill_color, outline=border_color, width=3)
        
    x1, y1, x2, y2 = coords
    cy = y1 + 15
    if title:
        draw.text((x1 + 15, cy), title, fill=DEEP_BLUE, font=font_title)
        cy += 25
        
    lines = text.split('\n')
    for line in lines:
        draw.text((x1 + 15, cy), line, fill=TEXT_DARK, font=font_text)
        cy += 18

# Helper function to draw arrow
def draw_arrow(draw, start, end, color=(113, 128, 150), width=2):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    arrow_size = 10
    a1 = angle + math.pi - math.pi / 6
    a2 = angle + math.pi + math.pi / 6
    ax1 = x2 + arrow_size * math.cos(a1)
    ay1 = y2 + arrow_size * math.sin(a1)
    ax2 = x2 + arrow_size * math.cos(a2)
    ay2 = y2 + arrow_size * math.sin(a2)
    draw.polygon([end, (ax1, ay1), (ax2, ay2)], fill=color)

# Generate System Architecture Image
def generate_system_architecture_image(path):
    img = Image.new('RGB', (1000, 500), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    
    # Load fonts
    try:
        font_title = ImageFont.truetype("arial.ttf", 16)
        font_text = ImageFont.truetype("arial.ttf", 12)
        font_header = ImageFont.truetype("arialbd.ttf", 20)
    except Exception:
        font_title = ImageFont.load_default()
        font_text = ImageFont.load_default()
        font_header = ImageFont.load_default()
        
    # Title
    draw.text((230, 20), "SYSTEM ARCHITECTURE & MEDIA FLOW PIPELINE", fill=DEEP_BLUE, font=font_header)
    
    # Boxes
    draw_box(draw, (50, 150, 320, 380), 
             "• Mobile & Web Layout\n• Custom UI Widgets\n• GoRouter Configurations\n• Provider State Management\n• Flutter WebRTC Call Plugin\n• Multi-Part Image Upload Service", 
             "1. FLUTTER CLIENT APP", font_title=font_title, font_text=font_text)
             
    draw_box(draw, (650, 60, 950, 230), 
             "• Session tokens & OAuth\n• PostgreSQL Relational Tables\n• RLS Security Policies\n• Real-Time Client Websockets\n• Supabase Storage Buckets", 
             "2. SUPABASE BACKEND", font_title=font_title, font_text=font_text)
             
    draw_box(draw, (650, 280, 950, 450), 
             "• YOLOv8 Bounding Box Drawing\n• Salesforce/BLIP Captioning\n• ViT Deepfake Classifier\n• Sentence-Transformers similarity\n• JSON API endpoint response", 
             "3. PYTHON FLASK AI SERVER", font_title=font_title, font_text=font_text)
             
    # Draw Arrows
    # Client to AI
    draw_arrow(draw, (320, 320), (650, 340), color=LIGHT_BLUE, width=3)
    draw.text((360, 295), "1. Upload Image (POST request)", fill=DEEP_BLUE, font=font_text)
    
    # AI to Client
    draw_arrow(draw, (650, 390), (320, 360), color=LIGHT_BLUE, width=3)
    draw.text((360, 370), "2. JSON (Bbox, Caption, Fake verdict)", fill=DEEP_BLUE, font=font_text)
    
    # Client to Supabase
    draw_arrow(draw, (320, 200), (650, 120), color=DEEP_BLUE, width=3)
    draw.text((370, 140), "3. Save Post & Media Buckets", fill=DEEP_BLUE, font=font_text)
    
    # Supabase to Client
    draw_arrow(draw, (650, 180), (320, 240), color=DEEP_BLUE, width=3)
    draw.text((370, 220), "4. WebSocket Real-time Feed Stream", fill=DEEP_BLUE, font=font_text)
    
    img.save(path)
    print(f"System Architecture diagram saved at: {path}")

# Generate Image Processing Pipeline Image
def generate_image_pipeline_image(path):
    img = Image.new('RGB', (1000, 620), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    
    # Load fonts
    try:
        font_title = ImageFont.truetype("arial.ttf", 15)
        font_text = ImageFont.truetype("arial.ttf", 11)
        font_header = ImageFont.truetype("arialbd.ttf", 18)
    except Exception:
        font_title = ImageFont.load_default()
        font_text = ImageFont.load_default()
        font_header = ImageFont.load_default()
        
    # Title
    draw.text((250, 20), "INTELLIGENT MULTI-MODEL IMAGE PROCESSING PIPELINE", fill=DEEP_BLUE, font=font_header)
    
    # Vertically arranged boxes
    draw_box(draw, (350, 60, 650, 110), "User snaps photo / selects media from gallery", "1. Image Input & Upload", font_title=font_title, font_text=font_text)
    draw_box(draw, (350, 140, 650, 190), "EXIF correction, auto-contrast, RGB conversion", "2. Image Preprocessing", font_title=font_title, font_text=font_text)
    
    # Split
    draw_box(draw, (40, 220, 310, 300), "Delineates objects, outputs label lists &\nconfidence scores (yolov8m.pt)", "3a. YOLOv8 Model", font_title=font_title, font_text=font_text)
    draw_box(draw, (350, 220, 650, 300), "Generates raw caption text describing\nmain elements (blip-large)", "3b. BLIP Caption Model", font_title=font_title, font_text=font_text)
    draw_box(draw, (690, 220, 960, 300), "Classifies media as Authentic Human or\nAI-Generated Synthetics", "3c. ViT Deepfake Classifier", font_title=font_title, font_text=font_text)
    
    # Merge
    draw_box(draw, (200, 340, 800, 420), 
             "NLTK tokenization + Sentence-Transformers (all-MiniLM-L6-v2) cosine similarity\nchecks. Compares BLIP caption terms against YOLO labels. Hallucinated missing objects\nare stripped from output caption text to ensure high semantic accuracy.", 
             "4. Semantic Cross-Check Layer", font_title=font_title, font_text=font_text)
             
    draw_box(draw, (350, 450, 650, 500), "Draws bounding boxes & converts to base64 string", "5. Bbox Annotation & Base64", font_title=font_title, font_text=font_text)
    draw_box(draw, (350, 530, 650, 580), "Uploads original to Supabase & returns JSON response", "6. DB Storage & Return JSON", font_title=font_title, font_text=font_text)
    
    # Connections
    draw_arrow(draw, (500, 110), (500, 140), color=LIGHT_BLUE, width=2)
    # Split arrows
    draw_arrow(draw, (500, 190), (175, 220), color=LIGHT_BLUE, width=2)
    draw_arrow(draw, (500, 190), (500, 220), color=LIGHT_BLUE, width=2)
    draw_arrow(draw, (500, 190), (825, 220), color=LIGHT_BLUE, width=2)
    # Merge arrows
    draw_arrow(draw, (175, 300), (500, 340), color=LIGHT_BLUE, width=2)
    draw_arrow(draw, (500, 300), (500, 340), color=LIGHT_BLUE, width=2)
    draw_arrow(draw, (825, 300), (500, 340), color=LIGHT_BLUE, width=2)
    
    draw_arrow(draw, (500, 420), (500, 450), color=LIGHT_BLUE, width=2)
    draw_arrow(draw, (500, 500), (500, 530), color=LIGHT_BLUE, width=2)
    
    img.save(path)
    print(f"Image Pipeline flowchart saved at: {path}")

# Generate WebRTC Flow Image
def generate_webrtc_flow_image(path):
    img = Image.new('RGB', (1000, 450), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    
    # Load fonts
    try:
        font_title = ImageFont.truetype("arial.ttf", 15)
        font_text = ImageFont.truetype("arial.ttf", 11)
        font_header = ImageFont.truetype("arialbd.ttf", 18)
    except Exception:
        font_title = ImageFont.load_default()
        font_text = ImageFont.load_default()
        font_header = ImageFont.load_default()
        
    # Title
    draw.text((260, 20), "WEBRTC REAL-TIME PEER-TO-PEER CALL SIGNALING FLOW", fill=DEEP_BLUE, font=font_header)
    
    # 3 Column boxes
    draw_box(draw, (50, 100, 300, 270), 
             "• Initiates Voice/Video Call\n• Generates SDP Offer payload\n• Gathers local ICE Candidate endpoints\n• Listens for dynamic signaling updates", 
             "CALLER DEVICE", font_title=font_title, font_text=font_text)
             
    draw_box(draw, (380, 100, 620, 270), 
             "• Relays signals in real-time\n• Custom Table: 'calls'\n• Exposes real-time stream subscription\n• Handles SDP and ICE handshakes", 
             "SUPABASE SIGNALING", font_title=font_title, font_text=font_text)
             
    draw_box(draw, (700, 100, 950, 270), 
             "• Receives incoming Call Offer\n• Generates SDP Answer payload\n• Gathers remote ICE Candidate endpoints\n• Establishes local device camera/mic", 
             "RECEIVER DEVICE", font_title=font_title, font_text=font_text)
             
    # Signaling arrows
    draw_arrow(draw, (300, 130), (380, 130), color=LIGHT_BLUE, width=2)
    draw.text((310, 115), "1. Offer SDP", fill=DEEP_BLUE, font=font_text)
    
    draw_arrow(draw, (620, 130), (700, 130), color=LIGHT_BLUE, width=2)
    draw.text((630, 115), "2. Stream Offer", fill=DEEP_BLUE, font=font_text)
    
    draw_arrow(draw, (700, 180), (620, 180), color=LIGHT_BLUE, width=2)
    draw.text((630, 165), "3. Answer SDP", fill=DEEP_BLUE, font=font_text)
    
    draw_arrow(draw, (380, 180), (300, 180), color=LIGHT_BLUE, width=2)
    draw.text((310, 165), "4. Stream Answer", fill=DEEP_BLUE, font=font_text)
    
    draw_arrow(draw, (300, 230), (380, 230), color=LIGHT_BLUE, width=2)
    draw_arrow(draw, (620, 230), (700, 230), color=LIGHT_BLUE, width=2)
    draw.text((450, 215), "5. Swap ICE Candidates", fill=DEEP_BLUE, font=font_text)
    
    # P2P link (bottom)
    draw.line([(175, 270), (175, 380), (825, 380), (825, 270)], fill=DEEP_BLUE, width=4)
    # Draw arrow heads at both ends
    draw_arrow(draw, (175, 280), (175, 270), color=DEEP_BLUE, width=4)
    draw_arrow(draw, (825, 280), (825, 270), color=DEEP_BLUE, width=4)
    draw.text((320, 390), "6. DIRECT PEER-TO-PEER WEBRTC MEDIA CHANNEL (AUDIO & VIDEO STREAMING)", fill=DEEP_BLUE, font=font_text)
    
    img.save(path)
    print(f"WebRTC signaling flowchart saved at: {path}")

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = docx.oxml.OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = docx.oxml.OxmlElement(f'w:{m}')
        node.set(docx.oxml.ns.qn('w:w'), str(val))
        node.set(docx.oxml.ns.qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(26, 54, 93) # Deep Blue
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(43, 108, 176) # Lighter Blue
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(74, 85, 104) # Slate Gray
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(4)
        
    heading.paragraph_format.keep_with_next = True
    return heading

def add_custom_paragraph(doc, text, space_after=6, line_spacing=1.15):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(45, 55, 72)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(44, 122, 123) # Dark Teal
    return p

def add_screenshot_placeholder(doc, label):
    """Add a beautifully styled placeholder table for screenshots."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading Accent 1'
    
    cell = table.rows[0].cells[0]
    set_cell_margins(cell, top=160, bottom=160, left=240, right=240)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    
    run = p.add_run(f"📷 [ INSERT {label.upper()} HERE ]")
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(43, 108, 176)
    
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(4)
    caption_p.paragraph_format.space_after = Pt(12)
    
    caption_run = caption_p.add_run(f"Fig: {label}")
    caption_run.font.name = 'Calibri'
    caption_run.font.italic = True
    caption_run.font.size = Pt(9.5)
    caption_run.font.color.rgb = RGBColor(113, 128, 150)
    
    return table

def main():
    # File locations
    img_arch = 'E:/Social_Media_App/system_architecture.png'
    img_pipe = 'E:/Social_Media_App/image_processing_pipeline.png'
    img_rtc  = 'E:/Social_Media_App/webrtc_signaling_flow.png'
    
    # Generate diagrams
    generate_system_architecture_image(img_arch)
    generate_image_pipeline_image(img_pipe)
    generate_webrtc_flow_image(img_rtc)
    
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ---------------------------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------------------------
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(60)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("SOCIAL MEDIA APPLICATION WITH MULTI-MODEL AI INTEGRATION AND REAL-TIME WEB RTC CALLS")
    title_run.font.name = 'Arial'
    title_run.font.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(26, 54, 93)
    title_p.paragraph_format.space_after = Pt(12)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("A Comprehensive Academic Project Report and Software Architecture Documentation")
    sub_run.font.name = 'Calibri'
    sub_run.font.italic = True
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(74, 85, 104)
    subtitle_p.paragraph_format.space_after = Pt(120)
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run("Submitted in partial fulfillment of the requirements for the degree of\nBachelor of Science in Computer Science / Software Engineering\n\nSubmitted by:\n")
    info_run.font.name = 'Calibri'
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(45, 55, 72)
    
    author_run = info_p.add_run("MUHAMMAD FARHAN\n")
    author_run.font.name = 'Arial'
    author_run.font.bold = True
    author_run.font.size = Pt(12)
    author_run.font.color.rgb = RGBColor(26, 54, 93)
    
    date_run = info_p.add_run("\nJune 2026")
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(10)
    
    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # TABLE OF CONTENTS PLACEHOLDER
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "TABLE OF CONTENTS", level=1)
    
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.line_spacing = 1.3
    
    toc_items = [
        ("Chapter 1: INTRODUCTION", "1"),
        ("  1.1 Project Overview & Motivation", "1"),
        ("  1.2 Objectives of the Project", "1"),
        ("  1.3 Scope of the Application", "2"),
        ("  1.4 High-Level Architecture Design", "2"),
        ("Chapter 2: LITERATURE SURVEY & TECHNOLOGY STACK", "3"),
        ("  2.1 Cross-Platform Mobile Development (Flutter & Dart)", "3"),
        ("  2.2 State Management and Navigation (Provider & GoRouter)", "3"),
        ("  2.3 Backend as a Service (Supabase & PostgreSQL)", "4"),
        ("  2.4 Real-time Communication (WebRTC & WebSocket)", "4"),
        ("  2.5 AI Pipeline Models (YOLOv8, BLIP, ViT-GPT2, AI-Image-Detector)", "4"),
        ("  2.6 Cloud Colab Deployment vs Local GPU Inference", "5"),
        ("Chapter 3: SYSTEM ANALYSIS & REQUIREMENT SPECIFICATION", "6"),
        ("  3.1 Functional Requirements", "6"),
        ("  3.2 Non-Functional Requirements", "6"),
        ("  3.3 Hardware and Software Specifications", "7"),
        ("Chapter 4: SYSTEM DESIGN & ARCHITECTURE", "8"),
        ("  4.1 System Architecture Diagram", "8"),
        ("  4.2 Image Processing Pipeline Flow", "9"),
        ("  4.3 WebRTC Peer-to-Peer Calling & Signaling Workflow", "10"),
        ("  4.4 Database Schema & Table Mapping", "11"),
        ("Chapter 5: IMPLEMENTATION & CODE STRUCTURE", "13"),
        ("  5.1 Project Folder Layout", "13"),
        ("  5.2 Database Table Creation SQL DDL", "14"),
        ("  5.3 Database Triggers & Functions", "15"),
        ("  5.4 Flutter Supabase Service Layer", "16"),
        ("  5.5 Python Flask AI Pipeline API (Colab & Local)", "17"),
        ("Chapter 6: USER INTERFACE FLOW & SNAPSHOTS", "19"),
        ("  6.1 Authentication Screens", "19"),
        ("  6.2 Home Feed & Story Operations", "19"),
        ("  6.3 AI Upload Pipeline Interface", "19"),
        ("  6.4 Chat & WebRTC Communication Screen", "20"),
        ("  6.5 Profile & Dynamic Settings Screen", "20"),
        ("Chapter 7: CONCLUSION & FUTURE ENHANCEMENTS", "21"),
        ("  7.1 Project Conclusion Summary", "21"),
        ("  7.2 Engineering Challenges Solved", "21"),
        ("  7.3 Future Scope & Enhancements", "21"),
        ("REFERENCES", "22")
    ]
    
    for title, pg in toc_items:
        leader = "." * (80 - len(title))
        run = toc_p.add_run(f"{title}{leader}{pg}\n")
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(74, 85, 104)

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # CHAPTER 1
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "Chapter 1: INTRODUCTION", level=1)
    
    add_styled_heading(doc, "1.1 Project Overview & Motivation", level=2)
    add_custom_paragraph(doc, 
        "Modern digital interaction revolves around social networking. Users expect quick, interactive, and intelligent "
        "media processing. Standard social media applications provide basic photo-sharing and texting features, but "
        "suffer from problems such as unverified or fake AI-generated media, manual tagging requirements, and restricted "
        "real-time communication options. This project introduces a highly customized, features-rich Social Media Application "
        "combining Cross-Platform Mobile Development, Cloud Backend Services, Real-Time Peer-to-Peer communications, and "
        "an Advanced Intelligent AI Processing Pipeline.")
    
    add_custom_paragraph(doc, 
        "The application provides users with standard social networking modules (authentication, customizable feeds, stories "
        "that automatically expire after 24 hours, likes, commenting, following) while layering edge-cutting tools. When "
        "a user uploads a photo, it undergoes deepfake verification, automatic object classification, and descriptive "
        "caption generation before insertion. Furthermore, real-time voice and video calls are natively integrated utilizing "
        "WebRTC, making this a unified, premium user-centric platform.")

    add_styled_heading(doc, "1.2 Objectives of the Project", level=2)
    add_custom_paragraph(doc, "The main objectives of this project are:")
    
    objectives = [
        "To develop a high-performance cross-platform application using Google's Flutter framework, ensuring native performance on both Android and iOS from a single codebase.",
        "To deploy a robust, secure relational database backend using Supabase, implementing Row Level Security (RLS) to safeguard user private data.",
        "To build a real-time, low-latency messaging module using Supabase Streams, enabling users to swap text messages instantaneously.",
        "To implement peer-to-peer audio and video calling natively inside the app using the WebRTC protocol, eliminating the need for paid third-party calling APIs.",
        "To create a Multi-Model Intelligent AI Pipeline using Python, running YOLOv8 for object detection, BLIP for image captioning, umm-maybe AI Image Detector for Deepfake detection, and Sentence-Transformers for semantic cross-verification.",
        "To manage storage buckets efficiently, uploading user avatars and post images via secure binary streams."
    ]
    for obj in objectives:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(obj)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(45, 55, 72)

    add_styled_heading(doc, "1.3 Scope of the Application", level=2)
    add_custom_paragraph(doc, 
        "The scope encompasses a complete mobile and web social ecosystem. The system handles secure email-based user signup "
        "and login, maintaining profile fields like username, bio, and avatar. On the feeds interface, users can create posts, "
        "delete their own posts, share others' posts onto their feed, react to posts using various reaction types, and write "
        "threaded comments. Users can upload statuses (stories) which are automatically cleared out by database-side "
        "scheduling. The WebRTC-based voice and video call system functions with room IDs allowing immediate peer connections.")

    add_styled_heading(doc, "1.4 High-Level Architecture Design", level=2)
    add_custom_paragraph(doc, 
        "The system follows a three-tier architecture: the client tier, backend server tier, and intelligent processing tier.")
    
    add_custom_paragraph(doc, 
        "1. Client Tier (Flutter Mobile App): Built using a Feature-First pattern, splitting code into cohesive directories "
        "like Auth, Feed, Profile, Chat, and Call features. The UI communicates with the Supabase API Client and the Flask AI API.")
    
    add_custom_paragraph(doc, 
        "2. Cloud Backend Tier (Supabase): Provides user session management, storage buckets for media assets, and a "
        "PostgreSQL database. Row Level Security (RLS) policies are active, and database triggers automate administrative tasks.")
    
    add_custom_paragraph(doc, 
        "3. Intelligent AI Tier (Flask Web Server): Hosts deep learning models. It receives images uploaded by clients, passes "
        "them through YOLOv8 and BLIP, resolves contradictions via NLTK/Sentence-Transformers, scores the authenticity via "
        "the ViT deepfake classification classifier, and returns base64-marked images and JSON metadata.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # CHAPTER 2
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "Chapter 2: LITERATURE SURVEY & TECHNOLOGY STACK", level=1)
    
    add_styled_heading(doc, "2.1 Cross-Platform Mobile Development (Flutter & Dart)", level=2)
    add_custom_paragraph(doc, 
        "Flutter is Google's UI toolkit designed to build beautiful, natively compiled applications for mobile, web, "
        "desktop, and embedded systems from a single codebase. Unlike other frameworks that rely on platform-specific native "
        "bridges, Flutter compiles directly to ARM machine code, yielding native performance. The application is built using "
        "Dart, an object-oriented, class-defined language optimized for client-side development. Dart's Ahead-Of-Time (AOT) "
        "compilation produces fast startups, while its Just-In-Time (JIT) compilation powers Flutter's famous Stateful Hot Reload, "
        "enabling rapid debugging.")

    add_styled_heading(doc, "2.2 State Management and Navigation (Provider & GoRouter)", level=2)
    add_custom_paragraph(doc, 
        "State management in Flutter is handled via the Provider package, a wrapper around InheritedWidget that makes state "
        "sharing and manipulation extremely straightforward. It ensures efficient rendering by updating only the widgets that "
        "explicitly consume modified data (e.g., dynamically updating the app-wide primary theme color or switching light/dark "
        "modes). Navigation is structured using GoRouter, a declarative routing package. GoRouter supports nested routes, deep "
        "linking, redirection rules based on authentication state, and uses stateful navigation shells to keep screen stacks alive "
        "during bottom-bar navigation switches.")

    add_styled_heading(doc, "2.3 Backend as a Service (Supabase & PostgreSQL)", level=2)
    add_custom_paragraph(doc, 
        "Supabase is an open-source Firebase alternative built around PostgreSQL, a powerful object-relational database. Unlike "
        "NoSQL solutions, PostgreSQL supports strict relationships, complex triggers, joins, and functions. Authentication "
        "is managed natively via Supabase Auth, which handles encrypted passwords, verification links, and secure JSON Web Tokens "
        "(JWT). Supabase Storage is utilized to upload raw media assets into discrete buckets (e.g., profiles, posts, stories). "
        "Database changes are sent to the client in real-time through WebSocket-based Supabase Streams.")

    add_styled_heading(doc, "2.4 Real-time Communication (WebRTC & WebSocket)", level=2)
    add_custom_paragraph(doc, 
        "WebRTC (Web Real-Time Communication) is a free, open-source project that provides mobile browsers and desktop applications "
        "with real-time communication (RTC) via simple APIs. It enables peer-to-peer audio and video streaming directly between client "
        "devices without routing media through a central server. Because clients cannot find each other's IP addresses due to firewalls "
        "and NAT, a WebSocket connection is used as a signaling channel to exchange Session Description Protocol (SDP) offers, answers, "
        "and Interactive Connectivity Establishment (ICE) candidates. This application leverages flutter_webrtc to integrate voice "
        "and video calling seamlessly.")

    add_styled_heading(doc, "2.5 AI Pipeline Models (YOLOv8, BLIP, ViT-GPT2, AI-Image-Detector)", level=2)
    add_custom_paragraph(doc, 
        "The intelligent image processing core is designed around a multi-model pipeline. Depending on the environment configurations, "
        "the following models are supported:")
    
    add_custom_paragraph(doc, 
        "1. YOLOv8 (You Only Look Once): An advanced object detection model. YOLOv8 Nano (yolov8n.pt) is deployed in resource-constrained "
        "environments (like Google Colab) to minimize inference latency, while YOLOv8 Medium (yolov8m.pt) runs on local GPU architectures "
        "for higher localization accuracy.")
    
    add_custom_paragraph(doc, 
        "2. Image Captioning Engines: Two primary models are supported. The 'nlpconnect/vit-gpt2-image-captioning' is a lightweight "
        "transformers model linking ViT encoder to GPT2 decoder, utilized in cloud Colab configurations. The 'Salesforce/blip-image-captioning-large' "
        "is a state-of-the-art BLIP model deployed for higher semantic description quality.")
    
    add_custom_paragraph(doc, 
        "3. Deepfake Detection Classifiers: Authenticity scoring is handled via Hugging Face pipeline classifiers. The model "
        "'umm-maybe/AI-image-detector' is utilized in both Colab and local server environments to determine whether an image is Real or AI-Generated.")

    add_styled_heading(doc, "2.6 Cloud Colab Deployment vs Local GPU Inference", level=2)
    add_custom_paragraph(doc, 
        "To make the project highly accessible and testable, the system supports dual backend deployment targets:")
    
    add_custom_paragraph(doc, 
        "• Google Colab Cloud Target: The Jupyter Notebook (AI_Models_Colab.ipynb) configures a free T4 GPU environment, downloads "
        "the model binaries, runs the Flask web server inside Colab, and exposes it using an Ngrok secure reverse tunnel "
        "('eatable-monument-gone.ngrok-free.dev'). This allows immediate testing by pointing the Flutter client's api_service "
        "to the tunnel domain without local Python installations.")
    
    add_custom_paragraph(doc, 
        "• Local GPU Target: Runs on a local CUDA-enabled GPU. Deploys the complete Python file pipeline including Sentence-Transformers "
        "('all-MiniLM-L6-v2') to verify term relationships between BLIP captions and YOLO detections, cleaning out hallucinations.")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # CHAPTER 3
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "Chapter 3: SYSTEM ANALYSIS & REQUIREMENT SPECIFICATION", level=1)
    
    add_styled_heading(doc, "3.1 Functional Requirements", level=2)
    add_custom_paragraph(doc, "The functional requirements define what the system must do:")
    
    funcs = [
        "Authentication: Users must be able to sign up using email, name, username, and password. The system must verify authentication state on launch and redirect non-authenticated users to the login screen.",
        "Profile Management: Authenticated users must be able to edit their full name, bio, and upload profile pictures, which are stored in the 'profiles' storage bucket.",
        "Social Interactions: Users must be able to create posts with photos, add captions, delete their posts, follow other users, unfollow users, react to posts, and comment on posts.",
        "Story System: Users must be able to publish stories (image or colored text background) which expire after 24 hours.",
        "Real-Time Chat: Users must be able to exchange messages in real-time. Message lists must dynamically update without manual refreshes.",
        "WebRTC Calls: Users must be able to initiate and receive voice/video calls. Signaling parameters must be shared securely via database channels.",
        "AI Image Verification: Every post image must be analyzed for objects, captioned, and evaluated for deepfake content. The system must overlay bounding boxes on the uploaded image."
    ]
    for fn in funcs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(fn)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(45, 55, 72)

    add_styled_heading(doc, "3.2 Non-Functional Requirements", level=2)
    add_custom_paragraph(doc, "The non-functional requirements define constraints and quality metrics:")
    
    non_funcs = [
        "Security: Row Level Security (RLS) must be enabled on all Supabase tables. Users can only write/edit their own posts, profiles, and messages, but can read public data.",
        "Performance: UI animations must maintain 60 FPS. API calls to the AI backend should complete within 3 seconds, displaying a clean loader in the client UI.",
        "Reliability: The real-time streams must reconnect automatically if the user's internet connection drops.",
        "Scalability: Database tables must utilize foreign key indexing to support quick queries as user counts grow."
    ]
    for nf in non_funcs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(nf)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(45, 55, 72)

    add_styled_heading(doc, "3.3 Hardware and Software Specifications", level=2)
    add_custom_paragraph(doc, "The specifications required to compile and run the project client and servers:")
    
    # Add Tables for Specifications
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Resource Type'
    hdr_cells[1].text = 'Minimum Specification'
    hdr_cells[2].text = 'Recommended Specification'
    for cell in hdr_cells:
        set_cell_margins(cell)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    specs = [
        ("Developer PC CPU", "Intel Core i5 (8th Gen) / AMD Ryzen 5", "Intel Core i7 / Apple M1/M2/M3"),
        ("Developer PC RAM", "8 GB DDR4", "16 GB or 32 GB DDR4/DDR5"),
        ("Graphics (for AI Models)", "Shared Intel HD / AMD Graphics", "NVIDIA Dedicated GPU (CUDA enabled, 6GB+ VRAM)"),
        ("Storage", "50 GB Free Hard Disk space", "SSD Storage with 100+ GB Free space"),
        ("Mobile Testing Device", "Android OS 8.0 / iOS 12", "Android OS 11+ / iOS 15+"),
        ("Operating System", "Windows 10 / macOS 11 Big Sur", "Windows 11 / macOS 14 Sonoma"),
        ("IDE / Compiler", "VS Code / Flutter SDK 3.10.x", "VS Code / Android Studio / Xcode / Flutter SDK 3.19.x+"),
        ("Python Version", "Python 3.8", "Python 3.10 / 3.11 with virtualenv")
    ]
    
    for item in specs:
        row = table.add_row()
        for idx, text in enumerate(item):
            cell = row.cells[idx]
            cell.text = text
            set_cell_margins(cell)
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(45, 55, 72)

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # CHAPTER 4
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "Chapter 4: SYSTEM DESIGN & ARCHITECTURE", level=1)
    
    add_styled_heading(doc, "4.1 System Architecture Diagram", level=2)
    add_custom_paragraph(doc, 
        "The system architecture details the network mappings and operational relationships between the mobile client UI, "
        "the cloud database layer, and the GPU-based machine learning server. The client layer acts as a visual interface, "
        "invoking HTTP/HTTPS REST calls to the AI backend and executing WebSocket connections to sync relational data automatically.")
    
    # Insert System Architecture Image
    doc.add_picture(img_arch, width=Inches(6.2))
    p_cap1 = doc.add_paragraph()
    p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap1 = p_cap1.add_run("Fig 4.1: System Architecture Diagram showing API, DB, and Real-Time connections")
    run_cap1.font.name = 'Calibri'
    run_cap1.font.italic = True
    run_cap1.font.size = Pt(9.5)
    run_cap1.font.color.rgb = RGBColor(113, 128, 150)
    p_cap1.paragraph_format.space_after = Pt(12)

    add_styled_heading(doc, "4.2 Image Processing Pipeline Flow", level=2)
    add_custom_paragraph(doc, 
        "Every posted photograph undergoes a 6-stage computer vision evaluation. Preprocessing stabilizes size, orientation, and "
        "brightness. YOLOv8 identifies objects and coordinates for local box markup drawing. BLIP-large generates textual descriptions. "
        "The semantic checking layer evaluates term similarity to strip hallucinations, and the deepfake classifier determines "
        "media authenticity before saving.")

    # Insert Pipeline Flow Image
    doc.add_picture(img_pipe, width=Inches(6.2))
    p_cap2 = doc.add_paragraph()
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap2 = p_cap2.add_run("Fig 4.2: Flowchart detailing Multi-Model Image Processing & Semantic Validation Pipeline")
    run_cap2.font.name = 'Calibri'
    run_cap2.font.italic = True
    run_cap2.font.size = Pt(9.5)
    run_cap2.font.color.rgb = RGBColor(113, 128, 150)
    p_cap2.paragraph_format.space_after = Pt(12)

    add_styled_heading(doc, "4.3 WebRTC Peer-to-Peer Calling & Signaling Workflow", level=2)
    add_custom_paragraph(doc, 
        "Rather than utilizing expensive voice servers, calling features run over direct peer channels. The system resolves "
        "NAT mappings by swapping Session Description Protocol (SDP) keys and ICE candidates. This messaging is synced via "
        "Supabase client streams, establishing direct audio/video communication.")

    # Insert WebRTC Image
    doc.add_picture(img_rtc, width=Inches(6.2))
    p_cap3 = doc.add_paragraph()
    p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap3 = p_cap3.add_run("Fig 4.3: WebRTC Call signaling sequence diagram utilizing Supabase streams")
    run_cap3.font.name = 'Calibri'
    run_cap3.font.italic = True
    run_cap3.font.size = Pt(9.5)
    run_cap3.font.color.rgb = RGBColor(113, 128, 150)
    p_cap3.paragraph_format.space_after = Pt(12)

    add_styled_heading(doc, "4.4 Database Schema & Table Mapping", level=2)
    add_custom_paragraph(doc, "Here is the exact description of the primary relational database tables:")
    
    tables_spec = [
        ("profiles", "Stores user profile configurations.", "id (UUID, PK, FK to auth.users), username (TEXT, Unique), full_name (TEXT), avatar_url (TEXT), bio (TEXT), created_at (TIMESTAMPTZ)"),
        ("posts", "Stores user uploaded posts.", "id (BIGINT, PK, Auto-Increment), user_id (UUID, FK to profiles), image_url (TEXT), caption (TEXT), objects_json (JSONB), deepfake_result (JSONB), created_at (TIMESTAMPTZ)"),
        ("shared_posts", "Handles posts shared to user feeds.", "id (BIGINT, PK), sharer_id (UUID, FK to profiles), original_post_id (BIGINT, FK to posts), original_user_id (UUID, FK to profiles), caption (TEXT), created_at (TIMESTAMPTZ)"),
        ("stories", "Stores disappearing status updates.", "id (BIGINT, PK), user_id (UUID, FK to profiles), image_url (TEXT), content (TEXT), bg_color (TEXT), original_post_id (BIGINT), expires_at (TIMESTAMPTZ), created_at (TIMESTAMPTZ)"),
        ("comments", "Stores comments made on posts.", "id (UUID, PK), post_id (BIGINT, FK to posts), user_id (UUID, FK to profiles), content (TEXT), created_at (TIMESTAMPTZ)"),
        ("reactions", "Stores post likes and emotions.", "id (BIGINT, PK), post_id (BIGINT, FK to posts), user_id (UUID, FK to profiles), reaction_type (TEXT), created_at (TIMESTAMPTZ)"),
        ("follows", "Manages the user follow-system.", "follower_id (UUID, PK, FK to profiles), following_id (UUID, PK, FK to profiles), created_at (TIMESTAMPTZ)"),
        ("messages", "Stores chat text records.", "id (BIGINT, PK), sender_id (UUID, FK to profiles), receiver_id (UUID, FK to profiles), text (TEXT), created_at (TIMESTAMPTZ)"),
        ("notifications", "Stores in-app notifications.", "id (BIGINT, PK), receiver_id (UUID, FK to profiles), sender_id (UUID, FK to profiles), type (TEXT), post_id (BIGINT), is_read (BOOLEAN), created_at (TIMESTAMPTZ)"),
        ("story_views", "Tracks story viewers.", "id (BIGINT, PK), story_id (BIGINT, FK to stories), viewer_id (UUID, FK to profiles), created_at (TIMESTAMPTZ)")
    ]
    
    for tbl_name, desc, fields in tables_spec:
        add_styled_heading(doc, f"Table: {tbl_name}", level=3)
        add_custom_paragraph(doc, f"Description: {desc}")
        add_custom_paragraph(doc, f"Columns: {fields}")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # CHAPTER 5
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "Chapter 5: IMPLEMENTATION & CODE STRUCTURE", level=1)
    
    add_styled_heading(doc, "5.1 Project Folder Layout", level=2)
    add_custom_paragraph(doc, 
        "The Flutter application uses a Feature-First modular structure. The project files are logically cataloged as follows:")
    
    add_code_block(doc, 
        "lib/\n"
        "├── core/\n"
        "│   ├── constants/       # App-wide string/numerical constants\n"
        "│   ├── router/          # AppRouter configuration (GoRouter)\n"
        "│   ├── secrets/         # Supabase API keys config\n"
        "│   ├── services/        # SupabaseClient and Call service wrappers\n"
        "│   ├── theme/           # App themes (light/dark mode colors)\n"
        "│   ├── utils/           # Time and text format helpers\n"
        "│   └── widgets/         # Shared global UI widgets\n"
        "└── features/\n"
        "    ├── auth/            # Login, Registration, Password Reset screens\n"
        "    ├── chat/            # ChatDetailScreen, ChatListScreen modules\n"
        "    ├── dashboard/       # Bottom navigation bar layout controller\n"
        "    ├── home/            # News Feed, Status bar, PostCard layouts\n"
        "    ├── notifications/   # In-app notifications listing page\n"
        "    ├── post_creation/   # Camera interfaces, Post upload workflows\n"
        "    ├── profile/         # Profile screen, follow listings\n"
        "    ├── search/          # Explore posts/users search page\n"
        "    ├── settings/        # Preferences and account settings\n"
        "    └── splash/          # Splash page loading animations\n")

    add_styled_heading(doc, "5.2 Database Table Creation SQL DDL", level=2)
    add_custom_paragraph(doc, 
        "The following SQL scripts initialize the relational database tables in the PostgreSQL engine under the public schema:")
    
    add_code_block(doc, 
        "-- Table: profiles\n"
        "CREATE TABLE public.profiles (\n"
        "    id uuid REFERENCES auth.users ON DELETE CASCADE PRIMARY KEY,\n"
        "    username text UNIQUE NOT NULL,\n"
        "    full_name text,\n"
        "    avatar_url text,\n"
        "    bio text,\n"
        "    created_at timestamp with time zone DEFAULT timezone('utc'::text, now()) NOT NULL\n"
        ");\n\n"
        "-- Table: posts\n"
        "CREATE TABLE public.posts (\n"
        "    id bigint GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,\n"
        "    user_id uuid REFERENCES public.profiles ON DELETE CASCADE NOT NULL,\n"
        "    image_url text NOT NULL,\n"
        "    caption text,\n"
        "    objects_json jsonb,\n"
        "    deepfake_result jsonb,\n"
        "    created_at timestamp with time zone DEFAULT timezone('utc'::text, now()) NOT NULL\n"
        ");\n\n"
        "-- Table: reactions\n"
        "CREATE TABLE public.reactions (\n"
        "    id bigint GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,\n"
        "    post_id bigint REFERENCES public.posts ON DELETE CASCADE NOT NULL,\n"
        "    user_id uuid REFERENCES public.profiles ON DELETE CASCADE NOT NULL,\n"
        "    reaction_type text NOT NULL,\n"
        "    created_at timestamp with time zone DEFAULT timezone('utc'::text, now()) NOT NULL,\n"
        "    UNIQUE (post_id, user_id)\n"
        ");")

    add_styled_heading(doc, "5.3 Database Triggers & Functions", level=2)
    add_custom_paragraph(doc, 
        "To prevent state errors and simplify user registration, a trigger is defined on the Supabase platform. "
        "When a new record is added to 'auth.users' via the signup screen, this trigger automatically creates the corresponding "
        "row in the 'public.profiles' table:")
    
    add_code_block(doc, 
        "-- Create Function to copy auth user to profiles\n"
        "CREATE OR REPLACE FUNCTION public.handle_new_user()\n"
        "RETURNS trigger AS $$\n"
        "BEGIN\n"
        "  INSERT INTO public.profiles (id, username, full_name, avatar_url)\n"
        "  VALUES (\n"
        "    new.id,\n"
        "    COALESCE(new.raw_user_meta_data->>'username', 'user_' || substr(new.id::text, 1, 8)),\n"
        "    COALESCE(new.raw_user_meta_data->>'full_name', ''),\n"
        "    new.raw_user_meta_data->>'avatar_url'\n"
        "  );\n"
        "  RETURN new;\n"
        "END;\n"
        "$$ LANGUAGE plpgsql SECURITY DEFINER;\n\n"
        "-- Create Trigger\n"
        "CREATE TRIGGER on_auth_user_created\n"
        "  AFTER INSERT ON auth.users\n"
        "  FOR EACH ROW EXECUTE FUNCTION public.handle_new_user();")

    add_styled_heading(doc, "5.4 Flutter Supabase Service Layer", level=2)
    add_custom_paragraph(doc, 
        "The 'SupabaseService' class in the Dart application encapsulates database mutations, storage uploads, and live streams:")
    
    add_code_block(doc, 
        "class SupabaseService {\n"
        "  static final client = Supabase.instance.client;\n\n"
        "  // Fetch stream of posts with profile metadata\n"
        "  static Stream<List<Map<String, dynamic>>> streamPosts() {\n"
        "    return client.from('posts_with_profiles')\n"
        "                 .stream(primaryKey: ['id'])\n"
        "                 .order('created_at', ascending: false);\n"
        "  }\n\n"
        "  // Upload raw image bytes to posts bucket\n"
        "  static Future<String> uploadImageBytes(Uint8List bytes, String bucket, String pathName) async {\n"
        "    await client.storage.from(bucket).uploadBinary(pathName, bytes);\n"
        "    return client.storage.from(bucket).getPublicUrl(pathName);\n"
        "  }\n\n"
        "  // React to a post\n"
        "  static Future<void> upsertReaction(dynamic postId, String reactionType) async {\n"
        "    final myId = client.auth.currentUser!.id;\n"
        "    await client.from('reactions').delete().eq('post_id', postId).eq('user_id', myId);\n"
        "    await client.from('reactions').insert({\n"
        "      'post_id': postId,\n"
        "      'user_id': myId,\n"
        "      'reaction_type': reactionType,\n"
        "    });\n"
        "  }\n"
        "}")

    add_styled_heading(doc, "5.5 Python Flask AI Pipeline API (Colab & Local)", level=2)
    add_custom_paragraph(doc, 
        "The python engine runs under one of two architectures. Below are both implementation codes:")
    
    add_styled_heading(doc, "Configuration A: Google Colab Testing Backend (Lightweight)", level=3)
    add_custom_paragraph(doc, 
        "The lightweight script run in Google Colab (AI_Models_Colab.ipynb) utilizes yolov8n.pt, vit-gpt2-image-captioning, "
        "and umm-maybe/AI-image-detector models. Exposed via Ngrok reverse tunnel:")
    
    add_code_block(doc, 
        "# Load Models\n"
        "obj_model = YOLO('yolov8n.pt')\n"
        "cap_model = VisionEncoderDecoderModel.from_pretrained('nlpconnect/vit-gpt2-image-captioning')\n"
        "cap_processor = ViTImageProcessor.from_pretrained('nlpconnect/vit-gpt2-image-captioning')\n"
        "cap_tokenizer = AutoTokenizer.from_pretrained('nlpconnect/vit-gpt2-image-captioning')\n"
        "fake_pipe = pipeline('image-classification', model='umm-maybe/AI-image-detector')\n\n"
        "@app.route('/process_all', methods=['POST'])\n"
        "def process_all():\n"
        "    file = request.files['image']\n"
        "    img = Image.open(file).convert('RGB')\n"
        "    draw = ImageDraw.Draw(img)\n\n"
        "    # 1. YOLOv8\n"
        "    obj_results = obj_model(img, verbose=False)\n"
        "    detections = []\n"
        "    for box in obj_results[0].boxes:\n"
        "        x1, y1, x2, y2 = box.xyxy[0].tolist()\n"
        "        name = obj_model.names[int(box.cls[0])]\n"
        "        draw.rectangle([x1, y1, x2, y2], outline='red', width=3)\n"
        "        detections.append({'object': name, 'confidence': float(box.conf[0])})\n\n"
        "    # 2. Caption\n"
        "    pixel_values = cap_processor(images=img, return_tensors='pt').pixel_values\n"
        "    output_ids = cap_model.generate(pixel_values)\n"
        "    caption = cap_tokenizer.decode(output_ids[0], skip_special_tokens=True)\n\n"
        "    # 3. Deepfake\n"
        "    fake_results = fake_pipe(img)\n"
        "    highest = max(fake_results, key=lambda x: x['score'])\n"
        "    is_real = 'human' in highest['label'].lower() or 'real' in highest['label'].lower()\n\n"
        "    buffered = io.BytesIO()\n"
        "    img.save(buffered, format='JPEG')\n"
        "    marked_image = base64.b64encode(buffered.getvalue()).decode()\n"
        "    return jsonify({'objects': detections, 'marked_image': marked_image, 'caption': caption, 'deepfake': {'is_real': is_real}})")

    add_styled_heading(doc, "Configuration B: Advanced Local Pipeline (with Semantic Check)", level=3)
    add_custom_paragraph(doc, 
        "Deploys higher quality models and includes a NLTK Tokenizer and Sentence-Transformers ('all-MiniLM-L6-v2') similarity "
        "check block to filter out caption hallucinations:")
    
    add_code_block(doc, 
        "obj_model = YOLO('yolov8m.pt')\n"
        "cap_processor = BlipProcessor.from_pretrained('Salesforce/blip-image-captioning-large')\n"
        "cap_model = BlipForConditionalGeneration.from_pretrained('Salesforce/blip-image-captioning-large')\n"
        "fake_pipe = pipeline('image-classification', model='umm-maybe/AI-image-detector')\n"
        "similarity_model = SentenceTransformer('all-MiniLM-L6-v2')\n\n"
        "def advanced_semantic_cross_check(raw_caption, detected_objects):\n"
        "    tokens = word_tokenize(raw_caption.lower())\n"
        "    yolo_embeddings = similarity_model.encode(detected_objects, convert_to_tensor=True)\n"
        "    invalid_words = []\n"
        "    for word in tokens:\n"
        "        word_emb = similarity_model.encode(word, convert_to_tensor=True)\n"
        "        score = util.cos_sim(word_emb, yolo_embeddings)\n"
        "        if torch.max(score).item() < 0.35:\n"
        "            invalid_words.append(word)\n"
        "    # Filter invalid words from raw_caption and return cleaned text...")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # CHAPTER 6
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "Chapter 6: USER INTERFACE FLOW & SNAPSHOTS", level=1)
    
    add_styled_heading(doc, "6.1 Authentication Screens", level=2)
    add_custom_paragraph(doc, 
        "The user authentication sequence is clean and responsive. On launching the app, the Splash Screen checks if "
        "a valid Supabase session token exists. If not, the application routes the user to the Login Screen. The Login Screen "
        "provides email input, secure password text field hiding, and redirect links to the Register Screen. The Register Screen "
        "collects additional information (username, full name) which is saved directly to the database via triggers upon authentication.")

    add_styled_heading(doc, "6.1.1 Splash Screen Interface", level=3)
    add_screenshot_placeholder(doc, "Application Splash Screen UI & Initializer flow")

    add_styled_heading(doc, "6.1.2 User Login Screen UI", level=3)
    add_screenshot_placeholder(doc, "Login Screen showing email, password inputs and redirect links")

    add_styled_heading(doc, "6.1.3 User Registration Screen UI", level=3)
    add_screenshot_placeholder(doc, "Register Screen collecting name, username, email, and password")

    add_styled_heading(doc, "6.2 Home Feed & Story Operations", level=2)
    add_custom_paragraph(doc, 
        "Once logged in, the user sees the Dashboard containing the Home Feed. At the top of the feed is a horizontal story/status "
        "bar. Clicking a story displays a full-screen image/text status with a progress indicator, auto-navigating to the next "
        "story after 5 seconds. Below the story bar is a vertical list of posts. Each post is rendered inside a PostCard containing "
        "the author's avatar, username, post image, caption, and reactions. Users can toggle reactions and expand the comments "
        "sheet to write real-time feedback.")

    add_styled_heading(doc, "6.2.1 Home News Feed Screen", level=3)
    add_screenshot_placeholder(doc, "Home Feed UI with user story bar and active posts feed list")

    add_styled_heading(doc, "6.2.2 Story (Status) View Screen", level=3)
    add_screenshot_placeholder(doc, "Disappearing Story view showing dynamic progress indicator bar")

    add_styled_heading(doc, "6.3 AI Upload Pipeline Interface", level=2)
    add_custom_paragraph(doc, 
        "Clicking the central floating action button triggers the Post Creation screen. Users can snap a photo or select an image "
        "from the phone's gallery. On selection, the image is previewed. Upon clicking upload, a loading screen displays while "
        "the raw image bytes are sent to the Flask AI server. The server processes the image and returns YOLO bounding boxes, "
        "the generated caption, and the Deepfake authenticity report. Bounding boxes are displayed overlaid on the image to "
        "indicate what objects were detected.")

    add_styled_heading(doc, "6.3.1 Post Creation and Media Selection", level=3)
    add_screenshot_placeholder(doc, "Post Creation Screen showing image selection preview and caption text input")

    add_styled_heading(doc, "6.3.2 YOLOv8 and Deepfake Classification Outputs", level=3)
    add_screenshot_placeholder(doc, "AI Upload Processing Screen with YOLO Bounding Boxes and Deepfake authenticity verdict")

    add_styled_heading(doc, "6.4 Chat & WebRTC Communication Screen", level=2)
    add_custom_paragraph(doc, 
        "Clicking the messaging icon opens the Chat List showing recent conversations. Tapping a conversation opens the "
        "ChatDetailScreen. Message bubbles are updated dynamically using Supabase real-time stream listeners. The chat window "
        "features voice and video call buttons. Pressing these buttons initializes a WebRTC room. The app shares connection "
        "offers/answers as specialized signaling entries in the Supabase database. The WebRTC calling page shows full-screen "
        "remote video feeds alongside local preview views.")

    add_styled_heading(doc, "6.4.1 Real-time Direct Messaging Screen", level=3)
    add_screenshot_placeholder(doc, "Chat detail conversation screen UI showing text messages exchange")

    add_styled_heading(doc, "6.4.2 WebRTC Audio/Video Call Screen", level=3)
    add_screenshot_placeholder(doc, "WebRTC Calling Screen showing remote peer video stream and local feed preview")

    add_styled_heading(doc, "6.5 Profile & Dynamic Settings Screen", level=2)
    add_custom_paragraph(doc, 
        "The application provides a dedicated User Profile showing the user's uploaded and shared posts in a grid. It also "
        "displays follower and following counts. The Settings Screen allows toggling dynamic themes and changing the primary "
        "theme color.")

    add_styled_heading(doc, "6.5.1 User Profile Screen UI", level=3)
    add_screenshot_placeholder(doc, "Profile Screen UI detailing bios, follower numbers, and post grid")

    add_styled_heading(doc, "6.5.2 App Settings & Dynamic Theme Screen", level=3)
    add_screenshot_placeholder(doc, "Settings Screen showing Dynamic Theme selection and customization panel")

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # CHAPTER 7
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "Chapter 7: CONCLUSION & FUTURE ENHANCEMENTS", level=1)
    
    add_styled_heading(doc, "7.1 Project Conclusion Summary", level=2)
    add_custom_paragraph(doc, 
        "The design and implementation of the Social Media Application successfully demonstrate the fusion of modern cross-platform "
        "mobile client programming, secure serverless cloud architectures, and deep-learning AI inference. By using Flutter, "
        "the client delivers desktop-like smooth interactions. Utilizing Supabase as the backend ensures enterprise-grade "
        "relational data management with minimal administrative overhead. The custom AI pipeline establishes a robust defense "
        "against synthetic media/deepfakes and improves user experience by generating natural text descriptions of photos automatically.")

    add_styled_heading(doc, "7.2 Engineering Challenges Solved", level=2)
    add_custom_paragraph(doc, 
        "During development, several engineering challenges were addressed:")
    
    add_custom_paragraph(doc, 
        "1. WebRTC Signaling: Deployed custom signaling tables inside Supabase. Using Supabase's live stream channels to exchange "
        "SDP parameters dynamically solved firewall traversal (STUN/TURN) mapping issues without needing dedicated signaling nodes.")
    
    add_custom_paragraph(doc, 
        "2. Semantic Mismatches in Caption Hallucinations: Captions generated by BLIP sometimes referenced items that did "
        "not exist. Creating a semantic checking layer with sentence-transformers resolved this, ensuring that generated captions "
        "contain only objects verified by YOLOv8.")

    add_styled_heading(doc, "7.3 Future Scope & Enhancements", level=2)
    add_custom_paragraph(doc, "The system can be enhanced in the future in the following ways:")
    
    futures = [
        "Group Video Calling: Scaling the WebRTC implementation to support multiparty calls using selective forwarding units (SFU).",
        "Generative AI Filters: Integrating stable diffusion engines in the upload pipeline to let users stylize photographs.",
        "Encrypted Direct Messaging: Implementing Signal Protocol-based end-to-end encryption (E2EE) for chats.",
        "Video Deepfake Analysis: Expanding the deepfake check pipeline to support real-time video frames analysis."
    ]
    for ft in futures:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(ft)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(45, 55, 72)

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # REFERENCES
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "REFERENCES", level=1)
    
    refs = [
        "[1] Google Developers, \"Flutter Documentation - Build Apps for Any Screen,\" 2026. [Online]. Available: https://flutter.dev/docs",
        "[2] Supabase Open Source Community, \"Supabase Documentation & Postgres Engine Reference Guide,\" 2026. [Online]. Available: https://supabase.com/docs",
        "[3] Real-Time Communications Working Group, \"WebRTC 1.0: Real-Time Communication Between Browsers,\" W3C Recommendation, 2021.",
        "[4] J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, \"You Only Look Once: Unified, Real-Time Object Detection,\" IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 2016.",
        "[5] J. Li, D. Li, C. Xiong, and S. Hoi, \"BLIP: Bootstrapping Language-Image Pre-training for Unified Image-Text Understanding and Generation,\" International Conference on Machine Learning (ICML), 2022.",
        "[6] N. Reimers and I. Gurevych, \"Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks,\" Empirical Methods in Natural Language Processing (EMNLP), 2019."
    ]
    
    ref_p = doc.add_paragraph()
    ref_p.paragraph_format.line_spacing = 1.3
    ref_p.paragraph_format.space_after = Pt(6)
    
    for r in refs:
        run = ref_p.add_run(f"{r}\n\n")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(74, 85, 104)

    # Save Document
    output_path = 'E:/Social_Media_App/Social_Media_App_Documentation.docx'
    try:
        doc.save(output_path)
        print(f"Word Document saved successfully at: {output_path}")
    except PermissionError:
        alt_path = 'E:/Social_Media_App/Social_Media_App_Documentation_v4.docx'
        try:
            doc.save(alt_path)
            print(f"Permission Denied (file probably open in Word). Saved successfully at alternative path: {alt_path}")
        except PermissionError:
            print("Error: Both paths are locked. Please close the document in Microsoft Word.")
            sys.exit(1)

if __name__ == '__main__':
    main()
